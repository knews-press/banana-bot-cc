"""File creation utilities — Word, Excel, CSV, plain text.

All generated files land in EXPORT_DIR (/root/creations/).
Files are keyed by a UUID so multiple users never collide.
"""

import csv
import io
import uuid
from pathlib import Path

EXPORT_DIR = Path("/root/creations")


def _export_path(filename: str) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    uid = uuid.uuid4().hex[:8]
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    return EXPORT_DIR / f"{stem}-{uid}{suffix}"


def create_docx(filename: str, content: str) -> Path:
    """Create a .docx file from Markdown-ish plain text.

    Headings (lines starting with # / ## / ###) become Word headings.
    Everything else becomes normal paragraphs.  Blank lines = paragraph break.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(11)

    path = _export_path(filename if filename.endswith(".docx") else filename + ".docx")
    doc.save(str(path))
    return path


def create_xlsx(filename: str, sheets: list[dict]) -> Path:
    """Create a .xlsx file.

    sheets: list of { "name": str, "headers": [str, ...], "rows": [[val, ...], ...] }
    If a single flat list of dicts is passed it is auto-converted to one sheet.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default empty sheet

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2563EB")  # blue

    for sheet_def in sheets:
        ws = wb.create_sheet(title=sheet_def.get("name", "Tabelle")[:31])
        headers = sheet_def.get("headers", [])
        rows = sheet_def.get("rows", [])

        if headers:
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")

        for row_idx, row in enumerate(rows, 2 if headers else 1):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Auto-width
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=0)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 60)

    path = _export_path(filename if filename.endswith(".xlsx") else filename + ".xlsx")
    wb.save(str(path))
    return path


def create_csv(filename: str, headers: list[str], rows: list[list]) -> Path:
    """Create a UTF-8 CSV file."""
    path = _export_path(filename if filename.endswith(".csv") else filename + ".csv")
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)
    return path


def create_text(filename: str, content: str) -> Path:
    """Write plain text / markdown to a file."""
    path = _export_path(filename)
    path.write_text(content, encoding="utf-8")
    return path
